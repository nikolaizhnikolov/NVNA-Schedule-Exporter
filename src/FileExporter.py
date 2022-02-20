from dataclasses import asdict
from itertools import count
from re import M
from unicodedata import name
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt
from openpyxl import Workbook, load_workbook
from openpyxl.worksheet.worksheet import Worksheet
from math import ceil
from collections import Counter

import ExporterLogger as logger
from DayData import Lecture
from ExporterUtil import resource_path, ExportTypes
from LectureDataForExport import GroupTypes, LectureDataForExport, LectureTypes
from OpenPyxlRowInsertOverride import insert_rows

file = Document()

def create_style(name, size, font_name, is_bold):
    styles = file.styles    
    cell_style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    cell_style.font.size = Pt(size)
    cell_style.font.name = font_name
    cell_style.font.bold = is_bold
    return cell_style

cell_style = create_style('Table Cells', 11, 'Calibri', False)
header_style = create_style('Table Header', 14, 'Calibri', True)


def export_simple_report(data, folder, file_name, file_type, weekly_indices):
    file_type = str(file_type)
    if(file_type.__eq__(ExportTypes.EXCEL)):
        return export_data_into_excel(data, folder, file_name, weekly_indices)
    elif(file_type.__eq__(ExportTypes.WORD)):
        return export_data_into_word(data, folder, file_name, weekly_indices)
    elif(file_type.__eq__(ExportTypes.PLAINTEXT)):
        return export_data_into_text(data, folder, file_name, weekly_indices)
    else:
        logger.info('File type not supported!\nRaising exception')
        raise ValueError('File type not supported!')
    
    
def format_lecture_data_for_monthly_report(data):
    lecture_list=[]
    # Extract all lecture names and group numbers
    for day in data:
        for lecture in day.lectures:
            lecture_list.append(LectureDataForExport(lecture))
    # Get unique values
    lecture_counter = Counter(lecture_list)
    for lecture in lecture_counter:
        lecture.occurences = lecture_counter[lecture]
    lecture_list = list(lecture_counter)
    lecture_list.sort(key=lambda l: (l.name_number, l.groups))
    # Open data file
    template_data = load_workbook(filename=resource_path('assets/data.xlsx'), data_only=True)
    # Get discipline number from name
    disciplines_sheet = template_data['disciplines']
    students_sheet = template_data['students']
    for lecture in lecture_list:
        for i in range (1, disciplines_sheet.max_row):
            cell_value = str(disciplines_sheet['B'+str(i)].value)
            if lecture.name.find(cell_value) != -1:
                lecture.name_number = disciplines_sheet['A'+str(i)].value
                break
        # Get student count and type from group number
        for group in lecture.groups:                
            for i in range (1, students_sheet.max_row):
                cell_value = str(students_sheet['A'+str(i)].value)
                if group.find(cell_value) != -1:
                    lecture.student_count += students_sheet['B'+str(i)].value
                    lecture.group_type = students_sheet['C'+str(i)].value
                    break    

    template_data.close()
    return lecture_list

def write_lecture_data_to_sheet(m_lectures, sheet, row_index):
    sheet.insert_rows = insert_rows
    for lecture in m_lectures:
        sheet.insert_rows(sheet, row_index, 1, above=False)
        # Get const index for each to know in which row to add
        sheet['A'+str(row_index)].value = lecture.name_number
        groups = ''
        # Set group number
        for group in lecture.groups:
            groups += group + ', '
        groups = groups[0:groups.__len__()-2]
        sheet['B'+str(row_index)].value = groups
        # Add student and lecture count
        if lecture.type == LectureTypes.LECTURE:
            sheet['C'+str(row_index)].value = lecture.student_count
            sheet['D'+str(row_index)].value = lecture.occurences
        # Add student count and practice count
        elif lecture.type == LectureTypes.PRACTICE:
            sheet['E'+str(row_index)].value = 1
            sheet['F'+str(row_index)].value = ceil(lecture.student_count / 2)
            sheet['G'+str(row_index)].value = lecture.occurences

        row_index+=1
        sheet.append(lecture.get_as_list())

    
def export_monthly_report(data, folder, file_name):
    # Load resources
    # Get all lecture data
    lecture_data = format_lecture_data_for_monthly_report(data)
    for lecture in lecture_data:
        logger.info(lecture.__hash__())

    file_name = file_name + ExportTypes.EXCEL
    file_path = folder + '\\' + file_name
    logger.info("Creating Excel Document: " + file_name + " in: " + folder)

    template = load_workbook(filename=resource_path('assets/template.xlsx'), data_only=False)
    sheet = template.active
    sheet.insert_rows = insert_rows

    M_ROW_INDEX = 42
    K_ROW_INDEX = 29
    B_ROW_INDEX = 14

    m_lectures = filter(lambda l: l.group_type == GroupTypes.MASTERS, lecture_data)
    write_lecture_data_to_sheet(m_lectures, sheet, 41)
    c_lectures = filter(lambda l: l.group_type == GroupTypes.CADET, lecture_data)
    write_lecture_data_to_sheet(c_lectures, sheet, 28)
    b_lectures = filter(lambda l: l.group_type == GroupTypes.BACHELORS, lecture_data)
    write_lecture_data_to_sheet(b_lectures, sheet, 13)

    
    template.save(file_path)
    template.close()

    return True    


def export_data_into_excel(data, folder, file_name, weekly_indices):
    file_name = file_name + ExportTypes.EXCEL
    file_path = folder + '\\' + file_name
    logger.info("Creating Excel Document: " + file_name + " in: " + folder)

    return True


def export_data_into_word(data, folder, file_name, weekly_indices):
    file_name = file_name + ExportTypes.WORD
    file_path = folder + '\\' + file_name
    logger.info("Creating Word Document: " + file_name + " in: " + folder)

    d_index = 0

    # For each week
    for w_index in weekly_indices:
        week_date_start = data[d_index].date
        week_date_end = data[d_index+6].date

        # Make new table
        table = file.add_table(rows=1, cols=5)
        table.style = 'TableGrid'
        table.autofit = False
        table.allow_autofit = False

        table_header = table.rows[0].cells
        
        # Set header info
        table_header[0].paragraphs[0].add_run('Седмица ' + str(w_index))
        table_header[1].paragraphs[0].add_run('От')
        table_header[2].paragraphs[0].add_run(str(week_date_start))
        table_header[3].paragraphs[0].add_run('До')
        table_header[4].paragraphs[0].add_run(str(week_date_end))

        # Get each day's info
        for i in range(7):
            day_info = data[d_index].info()
            if day_info is not None:   
                day_data = data[d_index]
                # Add header row and set data, then remove first lecture
                day_row = table.add_row().cells
                day_row[0].paragraphs[0].add_run(day_data.weekday).bold = True
                day_row[1].text = day_data.lectures[0].length
                day_row[2].text = day_data.lectures[0].lecture_name
                day_row[3].text = day_data.lectures[0].get_optionals()[0]
                day_row[4].text = day_data.lectures[0].get_optionals()[1]
                del(day_data.lectures[0])
                for lecture in day_data.lectures:
                    # Add new row for each other lecture
                    day_row = table.add_row().cells
                    day_row[1].text = lecture.length
                    day_row[2].text = lecture.lecture_name
                    day_row[3].text = lecture.get_optionals()[0]
                    day_row[4].text = lecture.get_optionals()[1]
            d_index+=1
        rows = table.rows.__len__()
        columns = table.columns.__len__()
        style_table(table, rows, columns)
            
    file.save(file_path)
    return True

def style_table(table, rows, columns):
    for column in range(columns):
        table.columns[column].width = Cm(get_column_width(column))

    for row in range(rows):
        row_cells = table.rows[row].cells
        for column in range(columns):
            if(row == 0):
                # Set a cell background (shading) color to dark grey.
                cell_color = parse_xml(r'<w:shd {} w:fill="A6A6A6"/>'.format(nsdecls('w')))
                row_cells[column]._tc.get_or_add_tcPr().append(cell_color)
                row_cells[column].paragraphs[0].style = header_style
            else:
                row_cells[column].paragraphs[0].style = cell_style
                row_cells[column].width = Cm(get_column_width(column))


def get_column_width(column_index):
    if (column_index == 0) | (column_index == 1):
        return 2.75
    elif (column_index == 2):
        return 5.5
    else:
        return 3.0


def export_data_into_text(data, folder, file_name, weekly_indices):
    file_name = file_name + ExportTypes.PLAINTEXT
    file_path = folder + '\\' + file_name
    
    logger.info("Creating Text File: " + file_name + " in: " + folder)

    # Format data and write into a file
    file = open(file_path, 'w', encoding='UTF-8')
    week_index_start = str(weekly_indices[0])
    week_index_end = str(weekly_indices[weekly_indices.__len__()-1])
    if week_index_start == week_index_end:
        file.write("Програма за седмица: " + week_index_start + '\n\n')
    else:
        file.write("Програма за седмици: " + week_index_start + " - " + week_index_end + '\n\n')
    # For each week
    d_index = 0
    for w_index in weekly_indices:
        week_date_start = data[d_index].date
        week_date_end = data[d_index+6].date
        if week_index_start != week_index_end:
            # Write the week index
            file.write('Седмица: ' + str(w_index) + '\n'+ \
                week_date_start + ' - ' + week_date_end + '\n')

        # Get each day's info
        for i in range(7):
            day_info = data[d_index].info()
            if day_info is not None:
                file.write(day_info + '\n')
            d_index+=1

    file.close()

    return True
